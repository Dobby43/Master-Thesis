from docx import Document
from docx.shared import Inches, Pt
import os


def build_report(output_path: str, text_data: dict):
    """
    Builds a Word document based on a given template (report.report_template.docx) by replacing keys marked as {key} inside the Word document with values or pictures (plots)

    :param output_path: Path to the generated .docx file
    :param text_data: Dictionary of text data with values to replace for the {keys}
    """
    # Load template from same folder as this function is in
    current_dir = os.path.dirname(os.path.abspath(__file__))
    template_path = os.path.join(current_dir, "report_template.docx")
    doc = Document(template_path)

    images_dict = {
        "table1": "gcode_plot.png",
        "table3": "characteristic_curve_plot.png",
    }

    # Go through the whole tables in report.report_template.docx and find keys which then get replaced
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    # replacement key from text_data
                    for key, value in text_data.items():
                        if f"{{{key}}}" in cell.text:
                            cell.text = cell.text.replace(f"{{{key}}}", str(value))

                    # Replace pictures marked with [[tableX]] in the report.report_template.docx
                    if "[[table1]]" in p.text and "table1" in images_dict:
                        p.clear()  # löscht alten Platzhaltertext
                        run = p.add_run()
                        run.add_picture(images_dict["table1"], width=Inches(6.0 / 2.54))

                    if "[[table3]]" in p.text and "table3" in images_dict:
                        p.clear()
                        run = p.add_run()
                        run.add_picture(images_dict["table3"], width=Inches(6.0 / 2.54))

    doc.save(output_path)
    print(f"\n[INFO] Report saved under {output_path}\n")


# Testfall
if __name__ == "__main__":
    pump_retract = True
    pump_control = True
    input = {
        "filename": "INPUT_FILE_STL",
        "volume": 0,
        "weight": 0,
        "measurements": [0, 0, 0],
        "x_min": 0,
        "y_min": 0,
        "z_min": 0,
        "x_max": 1,
        "y_max": 2,
        "z_max": 3,
        "robot_name": "KUKA KR340 R3330",
        "robot_geometry": {
            "a1": 500,
            "a2": 55,
            "b": 0,
            "c1": 1045,
            "c2": 1300,
            "c3": 1525,
            "c4": 290,
        },
        "tool_offset": {"X": -10.99, "Y": -0.86, "Z": 917.61},
        "tool_orientation": {"A": 0, "B": 0, "C": 180},
        "printbed": {"X": 1200, "Y": 4500, "Z": 2000},
        "base_coordinates": {
            "X": -1460.9,
            "Y": 2237.66,
            "Z": -268.5,
            "A": 0,
            "B": 0.0,
            "C": 0.0,
        },
        "print_speed": 0.35,
        "travel_speed": 0.5,
        "protract": -2,
        "retract": -1,
        "travel": 0,
        "outside": 1,
        "inside": 2,
        "surface": 3,
        "infill": 4,
        "bridge": 5,
        "curb": 6,
        "support": 7,
        "unknown": 99,
        "pump_retract": "activated" if pump_retract else "deactivated",
        "pump_control": "RPM" if pump_control else "Voltage",
        "flow_outer": 100,
        "flow_inner": 101,
        "flow_surface": 102,
        "flow_unknown": 103,
        "flow_infill": 104,
        "flow_bridge": 105,
        "flow_curb": 106,
    }

    build_report(output_path="test_report.docx", text_data=input)

    print(input["pump_retract"])
